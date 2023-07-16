import tkinter as tk
from openpyxl import Workbook, load_workbook
from tkinter import ttk
from tkinter import messagebox
import random
from tkinter import Tk, Label, Entry, Button, messagebox
from docx import Document
from openpyxl import load_workbook

class QuizManager:
    def __init__(self):
        self.filename = "questions.xlsx" # 問題資料的檔案名稱
        self.workbook = self.open_excel(self.filename) # 開啟或創建 Excel 檔案
        self.sheet_name = self.workbook.sheetnames[0] # 使用第一個工作表
        self.sheet = self.workbook[self.sheet_name] # 獲取工作表
        self.modified_data = [] # 儲存修改後的資料

        self.window = tk.Tk()
        self.window.title("題庫管理程式")  # 視窗標題

        self.num_questions_entry = None

        # 創建小工具
        self.create_widgets()

    # 初始化 Excel 檔案
    def open_excel(self, filename):
        try:
            workbook = load_workbook(filename)
            return workbook
        except FileNotFoundError:
            workbook = Workbook()
            workbook.create_sheet("Sheet1")
            del workbook["Sheet"]
            workbook.save(filename)
            return workbook

    # 創建小工具
    def create_widgets(self):
        self.question_label = ttk.Label(self.window, text="問題:")
        self.question_label.pack()
        self.question_entry = ttk.Entry(self.window)
        self.question_entry.pack(fill="x")

        self.answer_label = ttk.Label(self.window, text="答案:")
        self.answer_label.pack()
        self.answer_entry = ttk.Entry(self.window)
        self.answer_entry.pack(fill="x")

        self.submit_button = ttk.Button(self.window, text="新增", command=self.submit)
        self.submit_button.pack()

        self.modify_button = ttk.Button(self.window, text="修改", command=self.modify)
        self.modify_button.pack()

        # 總問題數量標籤
        num_questions = len(self.read_questions_from_excel(self.filename, self.sheet_name))
        self.total_label = Label(self.window, text='總問題數量：' + str(num_questions))
        self.total_label.pack()

        # 題目數量標籤和輸入框
        self.num_label = Label(self.window, text='題目數量：')
        self.num_label.pack()

        # 添加輸入框
        self.num_questions_entry = ttk.Entry(self.window)
        self.num_questions_entry.pack()
        self.num_questions_entry.insert(0, str(len(self.read_questions_from_excel(self.filename, self.sheet_name))))  # 預設顯示 5


        generate_button = ttk.Button(self.window, text="生成測驗", command=self.generate_quiz_and_export)
        generate_button.pack()


        self.status_label = ttk.Label(self.window, text="")
        self.status_label.pack()

        self.window.bind("<Return>", lambda event: self.submit())

    # 提交問題和答案
    def submit(self):
        question = self.question_entry.get()
        answer = self.answer_entry.get()

        self.add_question(question, answer)
        self.workbook.save(self.filename)

        preview_text = f"題目: {question}\n答案: {answer}"
        status_text = f"問題和答案已成功添加到工作表 '{self.sheet_name}'."
        self.status_label.config(text=status_text + "\n" + preview_text)

        num_questions = len(self.read_questions_from_excel(self.filename, self.sheet_name))
        self.total_label.config(text="總問題數量：" + str(num_questions))


        self.question_entry.focus()
        self.question_entry.delete(0, tk.END)
        self.answer_entry.delete(0, tk.END)

    # 將問題和答案添加到工作表
    def add_question(self, question, answer):
        if self.sheet.max_row == 1 and self.sheet.cell(row=1, column=1).value is None:
            next_row = 1
        else:
            next_row = self.sheet.max_row + 1

        self.sheet.cell(row=next_row, column=1, value=question)
        self.sheet.cell(row=next_row, column=2, value=answer)

    # 修改問題和答案
    def modify(self):
        modify_window = tk.Toplevel(self.window)
        modify_window.title("修改問題和答案")

        table = ttk.Treeview(modify_window, columns=("question", "answer"), show="headings")
        table.heading("question", text="問題")
        table.heading("answer", text="答案")

        rows = self.sheet.iter_rows(values_only=True)
        for row in rows:
            table.insert("", tk.END, values=row)

        table.pack()

        def edit_item(event):
            item_id = table.focus()
            column = table.identify_column(event.x)
            if column == "#1":
                entry_edit = tk.Entry(modify_window)
                entry_edit.insert(0, table.item(item_id)["values"][0])
                entry_edit.pack()
                entry_edit.focus_set()
                entry_edit.bind("<Return>", lambda event: save_edit(item_id, "#1", entry_edit.get()))
            elif column == "#2":
                entry_edit = tk.Entry(modify_window)
                entry_edit.insert(0, table.item(item_id)["values"][1])
                entry_edit.pack()
                entry_edit.focus_set()
                entry_edit.bind("<Return>", lambda event: save_edit(item_id, "#2", entry_edit.get()))

        def save_edit(item_id, column, edited_value):
            index = table.index(item_id)
            question = table.item(item_id)["values"][0]
            answer = table.item(item_id)["values"][1]
            # 擴展 self.modified_data 列表的長度
            while len(self.modified_data) <= index:
                self.modified_data.append(("",""))

            self.modified_data[index] = (question, answer)

            table.set(item_id, column=column, value=edited_value)

            self.sheet.delete_rows(1, self.sheet.max_row)
            for question, answer in self.modified_data:
                self.add_question(question, answer)

            self.workbook.save(self.filename)


        table.bind("<Double-Button-1>", edit_item)

        def delete_item():
            item_id = table.focus()
            if item_id:
                table.delete(item_id)

        def handle_delete(event):
            delete_item()

        delete_button = ttk.Button(modify_window, text="刪除", command=delete_item)
        delete_button.pack()
        modify_window.bind("<Delete>", handle_delete)

        def save_changes():
            self.modified_data = []
            for item_id in table.get_children():
                question = table.item(item_id)["values"][0]
                answer = table.item(item_id)["values"][1]
                self.modified_data.append((question, answer))

            existing_rows = self.sheet.max_row - 1
            self.sheet.delete_rows(1, self.sheet.max_row)

            for question, answer in self.modified_data:
                self.add_question(question, answer)

            self.workbook.save(self.filename)
            modify_window.destroy()
            self.question_entry.focus()
            self.status_label.config(text="問題和答案已成功修改.")

            num_questions = len(self.read_questions_from_excel(self.filename, self.sheet_name))
            self.total_label.config(text="總問題數量：" + str(num_questions))

        save_button = ttk.Button(modify_window, text="保存修改", command=save_changes)
        save_button.pack()

    # 從Excel檔案讀取問題
    def read_questions_from_excel(self, file_path, sheet_name):
        wb = load_workbook(file_path)
        sheet = wb[sheet_name]

        questions = []
        for row in sheet.iter_rows(values_only=True):
            question = {
                'question': row[0],
                'answer': row[1]
            }
            questions.append(question)

        return questions

    # 生成測驗
    def generate_quiz(self, questions, num_questions):
        quiz = random.sample(questions, num_questions)
        return quiz

    # 導出題目到Word檔案
    def export_to_docx(self, quiz, output_file):
        doc = Document()

        for i, question in enumerate(quiz, 1):
            doc.add_paragraph(f'問題 {i}:')
            doc.add_paragraph(str(question['question']))

        doc.save(output_file)

    # 導出答案到Word檔案
    def export_answers_to_docx(self, quiz, output_file):
        doc = Document()

        for i, question in enumerate(quiz, 1):
            doc.add_paragraph(f'{i}：')
            doc.add_paragraph(str(question['answer']))

        doc.save(output_file)




    # 生成測驗並導出到Word檔案
    def generate_quiz_and_export(self):
        num_questions = int(self.num_questions_entry.get())
        questions = self.read_questions_from_excel(self.filename, self.sheet_name)
        if len(questions) < num_questions:
            messagebox.showerror("錯誤", "工作表中的問題數量不足.")
            return

        quiz = self.generate_quiz(questions, num_questions)
            # 導出題目和答案到 Word 檔案

        try:
            self.export_to_docx(quiz, 'quiz.docx')
            self.export_answers_to_docx(quiz, 'answers.docx')
            messagebox.showinfo('成功', '題目word檔和答案word檔已成功導出。')
        except:
            messagebox.showerror('錯誤', '無法導出測驗和答案。')


    # 啟動程式
    def run(self):
        self.window.mainloop()


if __name__ == "__main__":
    quiz_manager = QuizManager()
    quiz_manager.run()
